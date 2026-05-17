from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_synopsis():
    doc = Document()

    # --- Title Page ---
    title = doc.add_heading('PROJECT SYNOPSIS', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph("\n" * 2)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("OSI Model Layer Visualization Tool")
    run.bold = True
    run.font.size = Pt(28)

    doc.add_paragraph("\n" * 3)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("Submitted By:\n").font.size = Pt(16)
    p.add_run("[Student Name]\n").font.size = Pt(18)
    doc.add_paragraph("\n")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("University: [University Name]\n").font.size = Pt(14)
    p.add_run("Department: Department of Computer Science & Engineering\n").font.size = Pt(14)
    p.add_run("Subject: Computer Networking and Communications\n").font.size = Pt(14)
    p.add_run("Project Guide: [Guide/Instructor Name]\n").font.size = Pt(14)
    p.add_run("Academic Year: 2023-2024\n").font.size = Pt(14)
    p.add_run("Submission Date: May 20, 2024").font.size = Pt(14)

    doc.add_page_break()

    # --- Objective ---
    doc.add_heading('1. Objective', level=1)
    objective_text = """
    The Open Systems Interconnection (OSI) model serves as the foundational architecture for modern telecommunications and computer networking. Developed by the International Organization for Standardization (ISO) in the 1980s, it partitions a communication system into seven logical layers, each with specific functions and protocols. This modular approach allows different manufacturers to develop hardware and software that can interoperate seamlessly. Understanding the OSI model is not merely an academic exercise; it is a critical skill for network engineers, cybersecurity analysts, and software developers. It provides a universal language for describing network operations and a structured methodology for diagnosing complex connectivity issues across heterogeneous environments.

    The primary objective of the OSI Model Layer Visualization Tool is to transform the abstract, often intimidating concepts of the 7-layer model into a tangible, interactive, and visually engaging experience. While traditional textbooks provide static diagrams and lengthy descriptions, they often fail to convey the dynamic nature of data transformation as it traverses the network stack. This tool provides a high-fidelity simulation of the encapsulation and decapsulation processes—the core mechanisms by which data is prepared for transmission and subsequently reconstructed by the receiver. By seeing these processes in action, students can better understand how each layer adds value and control information to the raw data, ensuring reliable and efficient communication.

    In real-world applications, the OSI model is indispensable for systematic troubleshooting. A network professional might start at the Physical layer to verify cable integrity, move to the Data Link layer to check for MAC address conflicts, and ascend to the Network layer to investigate routing table errors. This tool simulates these layers, allowing users to pause at any stage to inspect the Protocol Data Units (PDUs) like Segments, Packets, and Frames. Furthermore, simulation-based learning is proven to enhance cognitive retention by allowing students to experiment with different speeds and scenarios. This tool aims to be a comprehensive educational companion that bridges the gap between theoretical network models and the practical reality of data flow in a digital world.
    """
    doc.add_paragraph(objective_text)

    # --- Algorithms Explanation ---
    doc.add_heading('2. Algorithms and Concepts', level=1)

    # Encapsulation
    doc.add_heading('2.1 Encapsulation Process', level=2)
    encap_text = """
    Encapsulation is the fundamental algorithmic process by which data is wrapped in protocol-specific headers (and sometimes trailers) as it moves down the OSI stack from the Application layer to the Physical layer. Each layer takes the data unit from the layer above it, treats it as an opaque payload, and attaches its own control information. This control information is essential for the corresponding layer on the receiving system to understand how to process the data. Without encapsulation, the complex task of global networking would be unmanageable, as a single protocol would have to handle everything from user interface to electrical signaling.

    A classic real-world example of encapsulation is the process of sending a web request. The Application layer (HTTP) creates the request. The Presentation layer might encrypt this request using SSL/TLS. The Session layer manages the connection state. The Transport layer then breaks this data into segments and adds a TCP header for error recovery and flow control. The Network layer takes the segment and wraps it in an IP header containing source and destination IP addresses, creating a packet. The Data Link layer then wraps the packet in an Ethernet frame with MAC addresses and a Frame Check Sequence (FCS) for error detection. Finally, the Physical layer converts this frame into a series of electrical pulses or light signals.

    In the visualization tool, the encapsulation algorithm is implemented using a stack-based string manipulation approach. As the simulation moves down the layers, the system identifies the current active layer and retrieves its corresponding header metadata. This header is prepended to the existing payload string. The algorithm also updates the UI state, changing the color of the active layer's representation to provide immediate visual feedback. This step-by-step expansion of the packet demonstrates the "overhead" associated with networking, showing students that the data actually transmitted over the wire is significantly larger and more complex than the original message sent by the application.
    """
    doc.add_paragraph(encap_text)

    # Decapsulation
    doc.add_heading('2.2 Decapsulation Process', level=2)
    decap_text = """
    Decapsulation is the inverse algorithmic process of encapsulation, occurring at the receiving end of a network transmission. As the raw bits are received at the Physical layer, they are reassembled into frames, packets, segments, and finally, the original data. At each layer, the system examines the header intended for it, uses that information to manage the transmission (e.g., checking for errors or correct addressing), and then "peels off" or removes that header before passing the remaining payload up to the next higher layer in the stack. This ensures that the upper layers are never burdened with the technical details of the lower layers.

    Consider a web server receiving the aforementioned HTTP request. The Network Interface Card (NIC) at the Physical and Data Link layers receives the signal and strips the Ethernet header after verifying the MAC address and the FCS. The resulting IP packet is passed to the Network layer, which verifies the IP address and strips the IP header. The TCP segment is then passed to the Transport layer, which handles reassembly and stripping the TCP header. Finally, the Presentation and Session layers handle any decryption or session management before the raw HTTP request is delivered to the web server application. Each layer acts as a specialized filter, ensuring that only relevant data reaches the final destination.

    In our simulation, the decapsulation algorithm utilizes a robust parsing mechanism. As the packet moves up the receiver's stack, the algorithm identifies the outermost header block (e.g., "[Network_Header]"). It then removes this specific substring from the payload and updates the internal state of the packet object. The visualization reflects this by "shrinking" the displayed packet and highlighting the active receiver-side layer. This clearly demonstrates the recovery process, showing how the complex, encapsulated transmission unit is systematically simplified back into its original form, mirroring the high-efficiency processing that occurs in real-world network stacks.
    """
    doc.add_paragraph(decap_text)

    # Animation and Flow
    doc.add_heading('2.3 Animation and Data Flow Logic', level=2)
    anim_text = """
    The animation logic of the OSI Visualization Tool is governed by a sophisticated state machine that manages the 15 distinct stages of a complete end-to-end data transmission cycle. These stages include seven layers of downward movement (sender encapsulation), a central physical medium transition, and seven layers of upward movement (receiver decapsulation). The algorithm ensures that the flow is sequential and that the state of the packet and the UI are perfectly synchronized at every millisecond. This synchronization is crucial for providing an accurate educational representation of the deterministic nature of network protocols.

    The physical transmission step features a coordinate-based movement algorithm. Instead of a simple visibility toggle, the tool utilizes a Canvas-based animation where a visual "packet" object is moved across the screen using delta-X increments. The speed of this movement is dynamically calculated based on the user-defined speed slider, allowing for a smooth and continuous visual experience. This simulates the propagation delay inherent in physical media like fiber-optic cables or copper wires. The algorithm uses a non-blocking timing loop (via the after() method in Tkinter), which allows the UI to remain responsive to user inputs like "Pause" or "Reset" even while an animation is in progress.

    Data flow management also includes a robust "Pause and Resume" algorithm. When the user clicks "Pause," the state machine preserves all current variables—including the partially encapsulated payload, the current layer index, and the remaining animation frames. Upon "Resume," the algorithm calculates the remaining time and distance and continues the process from the exact point of interruption. This level of control is vital for educational purposes, as it allows instructors to stop the simulation at a critical juncture (like the Network layer) to explain specific concepts before allowing the data to continue its journey.
    """
    doc.add_paragraph(anim_text)

    # --- Pseudocode ---
    doc.add_heading('3. Pseudocode', level=1)

    doc.add_heading('3.1 Encapsulation and Downward Flow', level=2)
    doc.add_paragraph("PROCEDURE RunSenderSimulation(originalData):\n"
                      "    currentPacket = InitializePacket(originalData)\n"
                      "    FOR layerIndex FROM 0 TO 6:\n"
                      "        WHILE isPaused DO WAIT(100ms)\n"
                      "        IF isResetTriggered THEN EXIT PROCEDURE\n"
                      "        layerName = OSILayers[layerIndex].Name\n"
                      "        header = \"[\" + layerName + \"_Header]\"\n"
                      "        currentPacket.Payload = header + \" \" + currentPacket.Payload\n"
                      "        UI.HighlightSenderLayer(layerIndex)\n"
                      "        UI.UpdateInfoPanel(OSILayers[layerIndex].Description)\n"
                      "        UI.DisplayPayload(currentPacket.Payload)\n"
                      "        Delay(2000 / AnimationSpeed)\n"
                      "    TriggerPhysicalTransmission(currentPacket)\n"
                      "END PROCEDURE")
    doc.add_paragraph("The pseudocode for the sender-side simulation outlines a deterministic loop that iterates through the seven layers of the OSI model. It begins by creating a packet object with the user's initial data string. The loop then proceeds from the Application layer down to the Physical layer. Within each iteration, the algorithm checks for global flags such as 'paused' or 'reset'. The core logic involves string concatenation, where a layer-specific header is prepended to the payload. Simultaneously, the UI is updated to provide a visual highlight and display the educational text associated with that layer. The delay is inversely proportional to the user-controlled speed, ensuring that the simulation pace can be customized for different learning environments.")

    doc.add_heading('3.2 Physical Media Animation', level=2)
    doc.add_paragraph("PROCEDURE AnimatePhysicalMedia(packetElement):\n"
                      "    startX = 10, endX = 240, currentX = startX\n"
                      "    UI.ShowCanvas(True)\n"
                      "    WHILE currentX < endX:\n"
                      "        WHILE isPaused DO WAIT(100ms)\n"
                      "        IF isResetTriggered THEN EXIT PROCEDURE\n"
                      "        moveDistance = 5 * AnimationSpeed\n"
                      "        packetElement.PositionX += moveDistance\n"
                      "        currentX += moveDistance\n"
                      "        UI.RefreshCanvas()\n"
                      "        WAIT(20ms)\n"
                      "    UI.ShowCanvas(False)\n"
                      "    RunReceiverSimulation(currentPacket)\n"
                      "END PROCEDURE")
    doc.add_paragraph("The physical media animation pseudocode describes a coordinate-based movement logic. Once the packet has been fully encapsulated, it enters the transmission phase. The algorithm defines a starting and ending point on a graphical canvas. Using a high-frequency update loop (every 20ms), the packet's X-coordinate is incremented by a distance factor influenced by the simulation speed. This provides the 'smooth' movement requested in the project requirements. The loop includes checks for pause and reset states, ensuring a responsive user experience. Once the packet reaches the destination coordinate, the canvas is hidden, and the receiver-side decapsulation logic is triggered, passing the data to the next phase of the simulation.")

    doc.add_heading('3.3 Decapsulation and Upward Flow', level=2)
    doc.add_paragraph("PROCEDURE RunReceiverSimulation(encapsulatedPacket):\n"
                      "    FOR layerIndex FROM 0 TO 6:\n"
                      "        WHILE isPaused DO WAIT(100ms)\n"
                      "        IF isResetTriggered THEN EXIT PROCEDURE\n"
                      "        layer = OSILayers[6 - layerIndex] // Bottom-up\n"
                      "        headerToRemove = \"[\" + layer.Name + \"_Header]\"\n"
                      "        encapsulatedPacket.Payload.Remove(headerToRemove)\n"
                      "        UI.HighlightReceiverLayer(layerIndex)\n"
                      "        UI.UpdateInfoPanel(layer.Description)\n"
                      "        UI.DisplayPayload(encapsulatedPacket.Payload)\n"
                      "        Delay(2000 / AnimationSpeed)\n"
                      "    DisplaySuccessMessage(\"Data Received Successfully!\")\n"
                      "END PROCEDURE")
    doc.add_paragraph("Finally, the receiver-side pseudocode details the upward traversal of the OSI stack. This loop operates in reverse order compared to the sender, starting from the Physical layer and ending at the Application layer. The algorithm identifies the specific header string that was added during the encapsulation phase and removes it from the payload. This removal process is visually represented to show the packet 'shrinking' as it ascends the stack. Like the sender side, the UI is updated at each step with layer-specific educational content. Upon completion of the final iteration, a success message is displayed, confirming that the original data has been successfully recovered and delivered to the receiving application, completing the communication cycle.")

    # --- Conclusion ---
    doc.add_heading('4. Conclusion', level=1)
    conclusion_text = """
    The OSI Model Layer Visualization Tool represents a significant advancement in educational software for computer networking. By successfully implementing a dynamic, interactive simulation of the 7-layer OSI model, the project has achieved its primary goal of making abstract networking concepts accessible and understandable. The tool provides a comprehensive visual demonstration of encapsulation, decapsulation, and physical data transmission, allowing students to observe the intricate "handshakes" and data transformations that occur within a network stack. The use of modern Python libraries like CustomTkinter ensures that the tool is not only functional but also visually appealing and professional, meeting the standards expected for university-level submissions.

    Throughout the development of this project, several critical networking and software engineering principles were applied. The importance of modularity was reinforced through the separation of packet logic, animation management, and UI design. The deterministic nature of network protocols was modeled using a robust state machine, and the challenges of real-time UI synchronization were overcome using non-blocking asynchronous timing. These technical achievements mirror the very concepts the tool is designed to teach, such as layer independence and standardized communication interfaces.

    Looking ahead, there are numerous opportunities for further enhancement. Future versions of the tool could include support for simulating specific network protocols like IPv6, TCP vs. UDP behavior, and even modern concepts like Software-Defined Networking (SDN). Adding features to simulate network congestion, packet collision, and packet loss would provide students with a deeper understanding of error-correction mechanisms. Furthermore, expanding the tool into a multi-node simulation would allow for the visualization of routing and switching across complex topologies. In conclusion, this visualization tool is a powerful and versatile asset for any networking curriculum, providing a solid foundation for the next generation of network professionals.
    """
    doc.add_paragraph(conclusion_text)

    doc.save('synopsis.docx')
    print("synopsis.docx created successfully with expanded content.")

if __name__ == "__main__":
    create_synopsis()
